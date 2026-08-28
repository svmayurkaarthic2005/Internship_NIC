"""Pull plain text out of an uploaded file for the chat attachment store.

Supported: .txt, .csv, .pdf, .docx. The legacy binary .doc format is not
(no pure-Python reader is bundled) — callers should ask the user to save it as
.docx or PDF.
"""
from __future__ import annotations

import csv as _csv
import io as _io

try:
    from backend.config import settings as _settings
    MAX_DOC_PAGES = _settings.UPLOAD_MAX_DOC_PAGES
    EST_CHARS_PER_PAGE = _settings.UPLOAD_EST_CHARS_PER_PAGE
except Exception:  # keep the module importable outside the app (tests, tooling)
    MAX_DOC_PAGES = 15
    EST_CHARS_PER_PAGE = 1800

TEXT_EXTS = {".txt", ".csv"}
RICH_EXTS = {".pdf", ".docx"}
SUPPORTED_EXTS = TEXT_EXTS | RICH_EXTS

# Extensions we recognise but cannot read — message shown to the officer.
UNSUPPORTED_HINT = {
    ".doc": ("The legacy .doc format can't be read. Save it as .docx or PDF, "
             "or paste the text into the chat."),
}


class ExtractionError(Exception):
    """Raised when a recognised file type cannot be turned into text."""


def extract_text(ext: str, raw: bytes) -> str:
    ext = (ext or "").lower()
    if ext in TEXT_EXTS:
        return _from_text(ext, raw)
    if ext == ".pdf":
        return _from_pdf(raw)
    if ext == ".docx":
        return _from_docx(raw)
    raise ExtractionError(f"unsupported extension {ext!r}")


def _from_text(ext: str, raw: bytes) -> str:
    try:
        s = raw.decode("utf-8-sig")
    except UnicodeDecodeError:
        s = raw.decode("latin-1", errors="replace")
    if ext == ".csv":
        try:
            rows = list(_csv.DictReader(_io.StringIO(s)))
            if rows:
                return "\n".join(
                    " | ".join(f"{k}: {v}" for k, v in row.items())
                    for row in rows[:2000]
                )
        except Exception:
            pass  # not valid CSV — keep the raw text
    return s


def _from_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover - depends on environment
        raise ExtractionError("PDF support needs the 'pypdf' package installed.") from e
    try:
        reader = PdfReader(_io.BytesIO(raw))
        page_count = len(reader.pages)
    except Exception as e:
        raise ExtractionError(f"could not read the PDF ({e}).") from e

    if page_count > MAX_DOC_PAGES:
        raise ExtractionError(
            f"This PDF has {page_count} pages; the limit is {MAX_DOC_PAGES}. "
            f"Upload only the relevant pages (a long document overwhelms the "
            f"assistant and it starts to guess).")

    try:
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as e:
        raise ExtractionError(f"could not read the PDF ({e}).") from e
    body = "\n\n".join(f"[page {i + 1}]\n{t}" for i, t in enumerate(pages) if t)
    if not body.strip():
        raise ExtractionError(
            "No selectable text found in the PDF — it looks like a scan. "
            "Paste the text, or upload a text version.")
    return body


def _from_docx(raw: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover - depends on environment
        raise ExtractionError("Word support needs the 'python-docx' package installed.") from e
    try:
        document = docx.Document(_io.BytesIO(raw))
    except Exception as e:
        raise ExtractionError(f"could not read the Word file ({e}).") from e
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)

    # Word has no reliable page count without rendering — estimate from length.
    est_pages = max(1, -(-len(text) // EST_CHARS_PER_PAGE))  # ceil division
    if est_pages > MAX_DOC_PAGES:
        raise ExtractionError(
            f"This Word document is about {est_pages} pages of text; the limit "
            f"is {MAX_DOC_PAGES}. Upload a shorter extract (a long document "
            f"overwhelms the assistant and it starts to guess).")
    return text
